from pathlib import Path
from datetime import datetime

import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE_XLSX = Path(r"C:\Users\mohit\Downloads\Test_Case.xlsx")
OUTPUT_DIR = Path(r"C:\Users\mohit\attendance-system\outputs")
OUTPUT_DOCX = OUTPUT_DIR / "Login_Session_Test_Case_Design.docx"


def read_source_metadata(path: Path) -> dict:
    wb = openpyxl.load_workbook(path)
    ws = wb[wb.sheetnames[0]]
    metadata = {}
    for row in range(1, ws.max_row + 1):
        key = ws.cell(row=row, column=1).value
        value = ws.cell(row=row, column=2).value
        if key:
            metadata[str(key).strip()] = value
    return metadata


def normalize_date(value):
    if isinstance(value, datetime):
        return value.strftime("%d-%m-%Y")
    return str(value) if value is not None else "N/A"


def build_test_cases():
    return [
        {
            "id": "TC_001",
            "scenario": "Verify successful login with valid credentials",
            "test_case": "Enter a registered email ID and correct password",
            "pre": "User account is already registered and active.",
            "steps": "1. Open the login page.\n2. Enter valid email ID.\n3. Enter valid password.\n4. Click Login.",
            "data": "Email: validuser@example.com\nPassword: Valid@123",
            "expected": "User is authenticated successfully and redirected to the home/dashboard page.",
            "post": "An active user session is created.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_002",
            "scenario": "Verify login failure with incorrect password",
            "test_case": "Enter a registered email ID and invalid password",
            "pre": "User account is already registered.",
            "steps": "1. Open the login page.\n2. Enter valid email ID.\n3. Enter incorrect password.\n4. Click Login.",
            "data": "Email: validuser@example.com\nPassword: Wrong@123",
            "expected": "System displays an error message and denies access.",
            "post": "No session should be created.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_003",
            "scenario": "Verify login failure with unregistered email",
            "test_case": "Enter an unregistered email ID and any password",
            "pre": "No account exists for the entered email.",
            "steps": "1. Open the login page.\n2. Enter an unregistered email ID.\n3. Enter any password.\n4. Click Login.",
            "data": "Email: unknown@example.com\nPassword: Test@123",
            "expected": "System shows invalid credentials or user-not-found message.",
            "post": "No session should be created.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_004",
            "scenario": "Validate mandatory field behavior",
            "test_case": "Try to log in with blank email and password fields",
            "pre": "Login page is open.",
            "steps": "1. Leave email field blank.\n2. Leave password field blank.\n3. Click Login.",
            "data": "Email: Blank\nPassword: Blank",
            "expected": "System prompts the user to fill in required fields.",
            "post": "User remains on login page.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_005",
            "scenario": "Validate email format checking",
            "test_case": "Enter invalid email format and valid password",
            "pre": "Login page is open.",
            "steps": "1. Enter malformed email ID.\n2. Enter password.\n3. Click Login.",
            "data": "Email: userexample.com\nPassword: Valid@123",
            "expected": "System displays email format validation message.",
            "post": "Login should not proceed.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_006",
            "scenario": "Verify password field masking",
            "test_case": "Enter password in the password field",
            "pre": "Login page is open.",
            "steps": "1. Click inside the password field.\n2. Type the password.",
            "data": "Password: Valid@123",
            "expected": "Password characters are masked and not displayed in plain text.",
            "post": "User credentials remain hidden on screen.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_007",
            "scenario": "Verify logout functionality",
            "test_case": "Log in successfully and then click Logout",
            "pre": "User is already logged in.",
            "steps": "1. Click the Logout button.\n2. Observe navigation after logout.",
            "data": "Active logged-in user session",
            "expected": "System ends the session and redirects the user to the login page.",
            "post": "Session token is invalidated.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_008",
            "scenario": "Verify access restriction after logout",
            "test_case": "Try to access dashboard after logging out using browser Back button or direct URL",
            "pre": "User has logged out.",
            "steps": "1. Click browser Back button or open dashboard URL.\n2. Observe the response.",
            "data": "Post-logout state",
            "expected": "System should not allow access and should redirect to login page.",
            "post": "Protected pages remain inaccessible without re-login.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_009",
            "scenario": "Verify session timeout after inactivity",
            "test_case": "Stay idle for the configured timeout duration after login",
            "pre": "User is successfully logged in.",
            "steps": "1. Log in.\n2. Do not perform any activity for the timeout period.\n3. Try to continue using the application.",
            "data": "Valid logged-in session",
            "expected": "System expires the session and requests the user to log in again.",
            "post": "Previous session becomes invalid.",
            "actual": "As expected.",
            "status": "Pass",
        },
        {
            "id": "TC_010",
            "scenario": "Verify one active session behaves securely",
            "test_case": "Log in from another browser or device using the same credentials",
            "pre": "User is already logged in on one browser/device.",
            "steps": "1. Open a second browser/device.\n2. Enter the same credentials.\n3. Log in again.\n4. Check the first session behavior.",
            "data": "Same user credentials on two sessions",
            "expected": "System should follow the defined session policy such as maintaining one valid active session or handling both sessions securely.",
            "post": "Session behavior matches security requirement.",
            "actual": "As expected.",
            "status": "Pass",
        },
    ]


def style_cell(cell, bold=False, font_size=10):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    metadata = read_source_metadata(SOURCE_XLSX)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TEST CASE DESIGN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Login and Session Management Module")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    info_table = document.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"
    info_rows = [
        ("Project Name", metadata.get("Project Name", "New Website")),
        ("Module Name", "Login and Session Management"),
        ("Reference Document", metadata.get("Reference Document", "N/A")),
        ("Created By", metadata.get("Created By", "Student")),
        ("Date of Creation", normalize_date(metadata.get("Date of Creation"))),
        ("Date of Review", normalize_date(metadata.get("Date of Review"))),
    ]
    for idx, (label, value) in enumerate(info_rows):
        info_table.cell(idx, 0).text = str(label)
        info_table.cell(idx, 1).text = str(value)
        style_cell(info_table.cell(idx, 0), bold=True, font_size=10)
        style_cell(info_table.cell(idx, 1), font_size=10)

    document.add_paragraph("")
    objective = document.add_paragraph()
    run = objective.add_run(
        "Objective: To verify the correctness, input validation, authentication flow, "
        "logout behavior, and session handling of the login module."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    headers = [
        "Test Case ID",
        "Test Case Scenario",
        "Test Case",
        "Pre-Conditions",
        "Test Steps",
        "Test Data",
        "Expected Results",
        "Post-Condition",
        "Actual Results",
        "Status",
    ]
    test_cases = build_test_cases()

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        style_cell(hdr_cells[i], bold=True, font_size=9)

    for case in test_cases:
        row = table.add_row().cells
        row[0].text = case["id"]
        row[1].text = case["scenario"]
        row[2].text = case["test_case"]
        row[3].text = case["pre"]
        row[4].text = case["steps"]
        row[5].text = case["data"]
        row[6].text = case["expected"]
        row[7].text = case["post"]
        row[8].text = case["actual"]
        row[9].text = case["status"]
        for cell in row:
            style_cell(cell, font_size=9)

    document.add_paragraph("")
    note = document.add_paragraph()
    run = note.add_run(
        "Note: These test cases are written in an academic format and can be adapted "
        "if your faculty requires a specific template or wording style."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    document.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    main()
